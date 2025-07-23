from docx import Document
from io import BytesIO
from typing import Dict, Optional
import os
import json
from core.claude_client import get_claude_client

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_FOLDER = os.path.join(BASE_DIR, "../templates")


def autofill_placeholders(template_name: str, tender_result: dict) -> dict:
    client = get_claude_client()
    if not client:
        return {}  # fallback: empty dict

    prompt = f"""
    Ви — експерт із державних закупівель. Проаналізуйте наведені дані тендеру та згенеруйте JSON-словник значень для заповнення шаблону документа.

    Назва шаблону: {template_name}

    Витяг з аналізу тендеру:
    {json.dumps(tender_result, ensure_ascii=False, indent=2)}

    Поверніть тільки JSON-об'єкт, де ключі — це назви плейсхолдерів у шаблоні (наприклад: <Tender Title>), а значення — відповідні текстові вставки.
    """

    try:
        response = client.messages.create(
            model="claude-3-opus-20240229",
            max_tokens=500,
            temperature=0.2,
            messages=[{"role": "user", "content": prompt}]
        )
        return json.loads(response.content[0].text)
    except Exception as e:
        print("❌ Claude response parsing failed:", e)
        return {}


def generate_filled_template(template_filename: str, tender_result: Optional[Dict] = None) -> BytesIO:
    if not tender_result:
        raise ValueError("Missing tender_result")

    # 🧠 Generate values from Claude
    values = autofill_placeholders(template_filename, tender_result)

    # 📂 Load template DOCX
    path = os.path.join(TEMPLATE_FOLDER, f"{template_filename}.docx")
    if not os.path.exists(path):
        raise FileNotFoundError(f"Template not found: {path}")

    doc = Document(path)

    # 📝 Replace in paragraphs
    for p in doc.paragraphs:
        for run in p.runs:
            for key, val in values.items():
                if f"<{key}>" in run.text:
                    run.text = run.text.replace(f"<{key}>", val)

    # 📝 Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in values.items():
                    if f"<{key}>" in cell.text:
                        cell.text = cell.text.replace(f"<{key}>", val)

    # 📤 Return as BytesIO
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output
